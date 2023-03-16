from docx import Document

demo = Document()
for i in [1,2,3]:
    demo.add_heading('这是一个标题',level=1)
    demo.add_paragraph('这是一个段落')
    demo.save('demo.docx')

